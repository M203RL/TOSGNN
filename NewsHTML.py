##先將發布文章放在左螢幕的右半邊
import requests
from bs4 import BeautifulSoup
import urllib.request
from pathlib import Path
from urllib.parse import quote
import string
import os
import time
from docx import Document, oxml, opc
from docx.shared import Cm
from docx.oxml.ns import qn
from docx2pdf import convert
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import webbrowser
import pyperclip
import random
import pyimgur
import datetime
# from pynput.keyboard import Key, Listener
import pyautogui

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = oxml.shared.OxmlElement('w:r')
    rPr = oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def txt(list):
    try:
        text=''
        for item in list:
            text=text+item+'\n'
        pyperclip.copy(text)
        print('Text Copied')
    except TypeError:
        text=''
        for i in list:
            text=text+str(i)+'\n'
        pyperclip.copy(text)
        print('Text Copied')

def formatTime(time):
    if time<10:
        return '0'+str(time)
    else:
        return str(time)

kt='\'\\x16\''
def show(key):
  
    # print('\nYou Entered {0}'.format( key))
    # print(format( key))
    if format( key) == kt:
        # Stop listener
        return False

tStart=time.time()
t = time.localtime()

result = time.strftime("%Y/%m/%d %H:%M:%S", t)
print("Start Time: "+result)

CLIENT_ID = "886c33830062f60"
idx = 0

FILEBROWSER_PATH = os.path.join(os.getenv('WINDIR'), 'explorer.exe')
target = '慶祝活動'
cd = os.getcwd()
user_agent_list = ["Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/68.0.3440.106 Safari/537.36", "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/67.0.3396.99 Safari/537.36", "Mozilla/5.0 (Windows NT 10.0; WOW64) Gecko/20100101 Firefox/61.0","Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/64.0.3282.186 Safari/537.36","Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/62.0.3202.62 Safari/537.36","Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/45.0.2454.101 Safari/537.36","Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 6.0)","Mozilla/5.0 (Macintosh; U; PPC Mac OS X 10.5; en-US; rv:1.9.2.15) Gecko/20110303 Firefox/3.6.15"
                   ]

##test 場外測試
test=False
# test=True

##result 發布文章
result=False
result=True

##LinkSet 指定連結
LinkSet=False
# LinkSet=True

##reviewInput 心得
reviewInput=False
# reviewInput=True

##Alarm 發布提醒
Alarm=False
# Alarm=True

review='掉頭就走...'
# review='施工中...\n格式可能會亂掉'
if reviewInput:
    review=input("心得:")
print(review)


while True:
    
    url = 'https://towerofsaviors.com/category/%e5%85%ac%e5%91%8a/'
    
    headers = {'User-Agent': random.choice(user_agent_list)}
    response = requests.get(url=url, headers=headers)
    soup = BeautifulSoup(response.text, 'lxml')
    pArticle = soup.find_all('article')
    for i in range(10):
        h2 = pArticle[i].find('h2').text.strip()
        if target in h2:
            pA = pArticle[i].find('a')
            tPost=pArticle[i].find('time').get('datetime')
            try:
                timePost=tPost[tPost.index('T')+1:tPost.index('+')]
            except IndexError:
                timePost=tPost
            break
        
    if target in h2:
        ts=time.time()
        myLink = pA.get('href')
        

        if LinkSet:
            myLink = 'https://towerofsaviors.com/2022/10/07/%e3%80%90%e8%8b%b1%e9%9b%84%e7%8b%a9%e7%8d%b5%e9%96%8b%e5%a7%8b-%e2%80%a7-%e6%93%8a%e5%80%92%e9%9a%95%e7%9f%b3%e3%80%91%e6%85%b6%e7%a5%9d%e6%b4%bb%e5%8b%95/'

    
        newResponse = requests.get(url=myLink)
        newSoup = BeautifulSoup(newResponse.text, 'lxml')
        pArticle = newSoup.find('article')
        id = pArticle.get('id')
        


        folder = cd+'\\'+id
        latest = cd+'\\DL.txt'
        if not Path(latest).is_file():
            fi = open(latest, 'w')
            fi.close()
        fi = open(latest, 'r')
        rec = fi.read()
        if test:
            rec=''
        if id in rec:
            fi.close()
            pass
        else:
            print("\nNew Campaign Found")
            Path(folder).mkdir(parents=True, exist_ok=True)
            pThumbnail = pArticle.find('figure').find('img')
            pPhoto = pThumbnail['src']
            # print(pPhoto)
            s = quote(pPhoto, safe=string.printable)
            urllib.request.urlretrieve(s, folder+"\\cover.jpg")
            print('Thumbnail Downloaded')
            PATH = folder+"\\cover.jpg"
            im = pyimgur.Imgur(CLIENT_ID)
            uploaded_image = im.upload_image(PATH, title=id)

            pContent = pArticle.find('figure', {"class": "wp-block-table"})
            pNews = pContent.find('tbody')
            td = pNews.find_all('td')
            list = []
            doc = Document()
            doc.add_heading(h2, level=1)
            doc.add_picture(folder+"\\cover.jpg", width=Cm(15))
            tn='[div][img='+uploaded_image.link+' thumbnail=yes][/div]'
            list.append(tn)

            linkList=[]
            link=pNews.find_all('a')
            for links in link:
                ax=links['href']
                linkList.append(ax)
            uIndex=0

            for a in range(len(td)):
                af = str(td[a])
                af = str(td[a]).replace('<br/>', '\n')
                n1=af.find('<a href=')
                if n1!=-1:
                    n2=af.find('</a>')
                    st=af[n1:n2+4]
                    af=af.replace(st,'[url='+linkList[uIndex]+']'+linkList[uIndex]+'[/url]')
                    uIndex+=1
                af = af.replace('<strong>', '[b]')
                af = af.replace('</strong>', '[/b]')
                # af = af.replace(
                #     '<strong><mark class="has-inline-color has-carmack-yellow-color" style="background-color:rgba(0, 0, 0, 0)">', '[color=#790000]')
                af = af.replace(
                    '<mark class="has-inline-color has-carmack-yellow-color" style="background-color:rgba(0, 0, 0, 0)">', '[color=#790000]')
                af=af.replace('<mark class="has-inline-color" style="background-color:rgba(0, 0, 0, 0);color:#00ff0c">','')    
                af = af.replace('<td>', '[div][hr]')
                af = af.replace('</mark></strong>', '[/color]')
                af = af.replace('</mark>', '[/color]')
                af = af.replace('</td>', '[/div]')
                af = af.replace('&gt;', '>')
                list.append(af)
                # 
            text = ''
            for i in range(len(list)):
                text = text+list[i]
            # 
            text = text+'[hr][div]'+'[url='
            text = text+myLink
            text = text+'/]來源[/url] [/div]'
            text = text+'[div]更新時間: '+timePost+'[/div]'
            text = text+'[div]'+review+'[/div]'
            # tf=time.time()
            # dt=round(tf-ts,2)
            # text = text+'\nTotoal Time: '+str(dt)+'(+1.7) s'
            pyperclip.copy(text)
            if result:
                ts=time.time()
                pyperclip.copy(text)
                if Alarm:
                    webbrowser.open('https://www.youtube.com/watch?v=dQw4w9WgXcQ&t=0s',1)
                    os.startfile(folder+"\\cover.jpg")
                    input('Press Enter')
                #內文
                pyautogui.moveTo(-900, 600)
                pyautogui.click()
                # pyautogui.hotkey('ctrl', 'a')
                pyautogui.hotkey('ctrl', 'v')
                # with Listener(on_press = show) as listener:   
                #     listener.join()
                pyperclip.copy(h2)
                # print('Title Copied')

                # #標題
                pyautogui.moveTo(-900, 410)
                pyautogui.click()
                # # pyautogui.hotkey('ctrl', 'a')
                pyautogui.hotkey('ctrl', 'v')

                #預覽
                pyautogui.moveTo(-180, 170)
                pyautogui.click()

                #關閉預覽
                pyautogui.moveTo(-42, 168)
                pyautogui.click()

                #發佈
                pyautogui.moveTo(-100, 170)
                pyautogui.click()
                #確認
                pyautogui.moveTo(-530, 760)
                pyautogui.click()
                # print('Article Copied')
                # subprocess.run([FILEBROWSER_PATH, folder])
                tf=time.time()
                dt=round(tf-ts,2)
                print('Totoal Time: '+str(dt)+'s')
                


                print('Making PDF...')
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "思源黑體 HW"
                        r = run._element.rPr.rFonts  # 中文字型
                        r.set(qn('w:eastAsia'), "思源黑體 HW")
                doc.save(folder+"\\News.docx")
                convert(folder+"\\News.docx")
            
            if not test:
                fi = open(latest, 'w')
                fi.write(id+'\n'+rec)
                fi.close()
                break

            
            break
        tCurrent=time.time()
        tLapsed=round(tCurrent-tStart)
        m, s = divmod(tLapsed, 60)
        h, m = divmod(m, 60)
        h=formatTime(h)
        m=formatTime(m)
        s=formatTime(s)
        print('Time Lapsed: '+h+':'+m+':'+s+', Loops: '+str(idx), end="\r")
        delay_choices = [8, 5, 10, 6, 11, 9, 3,1,2,7]  # 延遲的秒數
        delay = random.choice(delay_choices)  # 隨機選取秒數
        time.sleep(delay)
        idx += 1
